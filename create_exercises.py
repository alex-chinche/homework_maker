from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive
import os
from docx import Document
import datetime
from docx.shared import Pt
from docx.shared import RGBColor
from colorama import init, Fore
init()

# Before get started, you need to enable google drive api here:
# https://www.google.com/url?sa=t&rct=j&q=&esrc=s&source=web&cd=&cad=rja&uact=8&ved=2ahUKEwjb2bTfqpvuAhXRRBUIHYaiAloQFjAAegQIBRAC&url=https%3A%2F%2Fdevelopers.google.com%2Fdrive%2Fapi%2Fv3%2Fenable-drive-api&usg=AOvVaw0_PTvjxowwsoESoTqMmPA3
# After that, make sure to download the OAuth json credentials to the same directory than this file, and rename that json to "client_secrets.json"

###########################
####### VARIABLES #########
###########################

DOCX_SAVE_PATH = os.getcwd() + "/exercises/"

###########################
####### FUNCTIONS #########
###########################


def style_question_title(*args):
    my_title = document.add_paragraph()
    my_title_style = my_title.add_run(
        question_title_text + str(number_of_exercise))
    my_title_style.bold = True
    my_title_style.font.color.rgb = RGBColor(25, 25, 112)
    my_title_style.font.size = Pt(20)
    my_title_style.underline = True
    my_title_style.font.name = 'Chilanka'
    return my_title


def style_answer_title(*args):
    my_title = document.add_paragraph()
    my_title_style = my_title.add_run(
        answer_title_text)
    my_title_style.bold = True
    my_title_style.font.color.rgb = RGBColor(25, 25, 112)
    my_title_style.font.size = Pt(20)
    my_title_style.underline = True
    my_title_style.font.name = 'Chilanka'
    return my_title


def style_paragraph(paragraph_text):
    my_paragraph = document.add_paragraph()
    my_paragraph_style = my_paragraph.add_run(paragraph_text)
    my_paragraph_style.font.size = Pt(15)
    my_paragraph_style.font.name = 'Chilanka'
    return my_paragraph


###################################
########## SCRIPT START ##########
###################################

today = datetime.datetime.now()
document_title = today.strftime('%B_%d_%Y')
document = Document()
page_heading = document.add_heading(0)
my_page_heading_style = page_heading.add_run(
    'Ejercicios ' + today.strftime('%B/%d/%Y') + '\n')
my_page_heading_style.font.size = Pt(32)
my_page_heading_style.font.name = 'Chilanka'
number_of_exercise = 0
question_title_text = 'Pregunta nº '
answer_title_text = 'Respuesta:'
# Loop start
welcome_message = input(
    Fore.YELLOW + "Hi Alex! Ready to give Diana some homework?(y/n)\n" + Fore.WHITE)
if welcome_message == "y":
    while True:
        number_of_exercise = number_of_exercise+1
        style_question_title(question_title_text, number_of_exercise)
        question_text = input(Fore.YELLOW + "Write the question number " +
                              str(number_of_exercise) + ':\n' + Fore.WHITE)
        to_remember = '"En el camino de la programación y la informática, habrá infinidad de momentos en los cuales nos sentiremos perdidos y pensaremos que es demasiado difícil, que no valemos para esto... Es esa sensación la que tenemos que vencer para superarnos día a día. Recuérdalo"'
        style_paragraph(question_text)
        style_answer_title(answer_title_text)
        document.add_paragraph("-")
        document.add_paragraph(
            "__________________________________________________________________________")
        print("...")
        print(Fore.GREEN + "Exercise created!")
        continue_confirmation = input(
            Fore.YELLOW + "Want to create more exercises? (Already " + str(number_of_exercise) + ")(y/n):\n" + Fore.WHITE)
        if continue_confirmation == "n":
            style_paragraph(to_remember)
            ### Saves the .docx ###
            my_file_path = DOCX_SAVE_PATH + document_title + '.docx'
            document.save(my_file_path)
            print("...")
            print(Fore.GREEN + "Ok. Check your folder. You have the document there!")
            upload_confirmation = input(Fore.YELLOW +
                                        "Do you want to upload the file to Google Drive?(y/n)\n" + Fore.WHITE)
            if upload_confirmation == "y":
                # Google Drive upload
                # Authentication
                gauth = GoogleAuth()
                drive = GoogleDrive(gauth)
                created_folder = drive.CreateFile({'title': today.strftime(
                    '%B_%d_%Y'), 'mimeType': 'application/vnd.google-apps.folder', 'parents': [{'id': '1iouWY37SFXUUNT0Rg1ZoRNAGGJB-_mrI'}]})
                print(Fore.YELLOW + "Creating folder...")
                created_folder.Upload()
                print(Fore.GREEN + "Folder created!")
                print(Fore.YELLOW + "Uploading file to drive...")

                created_doc_file = drive.CreateFile(
                    {'title': "Ejercicios " + today.strftime(
                        '%B_%d_%Y'), 'parents': [{'id': created_folder.get('id')}], 'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'})
                created_doc_file.SetContentFile(
                    DOCX_SAVE_PATH + document_title + '.docx')
                created_doc_file.Upload(param={'convert': True})

                print(Fore.GREEN + "File uploaded and everything done!")
            else:
                pass
            print(Fore.YELLOW + "Have a good day. Bye!!")
            break
        else:
            pass
else:
    print(Fore.YELLOW + "Ok, see you next time!!")
